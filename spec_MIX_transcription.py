# Skript um gezielt spezifische multilinguale Audiodateien zu transkribieren 


import os
import whisper
import torch
from docx import Document
import numpy as np

# Prüfen, ob eine GPU verfügbar ist
# an dieser Stelle sinnvoll, da mit dem large-Modell gearbeitet wird
device = "cuda" if torch.cuda.is_available() else "cpu"
if device == "cuda":
    print(f"GPU wird verwendet: {torch.cuda.get_device_name(0)}")
else:
    print("Keine GPU verfügbar. Whisper läuft auf der CPU.")

# Ordnerpfade
input_folder = "Dream recordings"
output_folder = "Dream transcripts"

# Whisper-Modell laden (large für höhere Genauigkeit)
model = whisper.load_model("large", device=device)

def transcribe_audio_chunk(chunk):
    """Transkribiert einen Chunk einer Audiodatei mit automatischer Spracherkennung."""
    result = model.transcribe(chunk, language=None, task="transcribe")  # Automatische Spracherkennung
    return result["text"]

def split_audio_into_chunks(audio_path, chunk_duration=10):
    """Teilt die Audiodatei in Chunks (Standard: 10 Sekunden)."""
    audio = whisper.load_audio(audio_path)
    chunk_samples = int(chunk_duration * whisper.audio.SAMPLE_RATE)

    # Teile die Audiodatei in Chunks
    return [audio[i:i + chunk_samples] for i in range(0, len(audio), chunk_samples)]

# Sicherstellen, dass der Ausgabeordner existiert
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 📝 Dateiname manuell eingeben
file_name = input("Bitte den Namen der Audiodatei (inkl. Endung, z.B. 'audio.mp3') eingeben: ")
audio_path = os.path.join(input_folder, file_name)

# Prüfen, ob die Datei existiert
if not os.path.isfile(audio_path):
    print(f"Fehler: Die Datei '{file_name}' wurde im Ordner '{input_folder}' nicht gefunden.")
else:
    try:
        print(f"Verarbeite Datei: {file_name}")

        # Audiodatei in 10-Sekunden-Chunks teilen
        audio_chunks = split_audio_into_chunks(audio_path, chunk_duration=10)

        # Word-Dokument erstellen
        doc = Document()
        doc.add_heading(f'Transkript der Datei: {file_name}', 0)

        # Transkription der Chunks
        for idx, chunk in enumerate(audio_chunks):
            print(f"Verarbeite Chunk {idx + 1} von {len(audio_chunks)}...")
            text = transcribe_audio_chunk(chunk)

            # Jeden Satz in eine eigene Zeile schreiben
            sentences = text.replace("?", "?.").replace("!", "!.").split(". ")
            for sentence in sentences:
                if sentence.strip():
                    doc.add_paragraph(sentence.strip())

        # Speichern der Datei
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}.docx")
        doc.save(output_file_path)

        print(f"✅ Transkription erfolgreich gespeichert unter: {output_file_path}")

    except Exception as e:
        print(f"❌ Ein Fehler ist aufgetreten: {e}")
