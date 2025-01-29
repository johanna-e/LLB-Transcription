# Multilinguale Transkription

import os
import whisper
import torch
from docx import Document
import numpy as np

# Prüfen, ob eine GPU verfügbar ist
device = "cuda" if torch.cuda.is_available() else "cpu"
if device == "cuda":
    print(f"GPU wird verwendet: {torch.cuda.get_device_name(0)}")
else:
    print("Keine GPU verfügbar. Whisper läuft auf der CPU.")

# Ordnerpfade
input_folder = "Dream recordings"
output_folder = "Dream transcripts"

# Whisper-Modell laden (turbo oder medium empfohlen)
model = whisper.load_model("large", device=device)

def get_latest_audio_file(folder_path):
    """Gibt die neueste Audiodatei im Ordner zurück."""
    audio_files = [f for f in os.listdir(folder_path) if f.endswith(('.mp3', '.wav', '.m4a'))]
    if not audio_files:
        return None
    return max(audio_files, key=lambda f: os.path.getmtime(os.path.join(folder_path, f)))

def transcribe_audio_chunk(chunk):
    """Transkribiert einen Chunk einer Audiodatei mit automatischer Spracherkennung."""
    result = model.transcribe(chunk, language=None, task="transcribe")  # Automatische Spracherkennung
    text = result["text"]
    language = result["language"]
    return text, language

def split_audio_into_chunks(audio_path, chunk_duration=30):
    """Teilt die Audiodatei in Chunks (Standard: 30 Sekunden)."""
    audio = whisper.load_audio(audio_path)
    audio_length = len(audio) / whisper.audio.SAMPLE_RATE  # Dauer in Sekunden
    chunk_samples = int(chunk_duration * whisper.audio.SAMPLE_RATE)

    # Teile die Audiodatei in Chunks
    chunks = []
    for start_sample in range(0, len(audio), chunk_samples):
        chunk = audio[start_sample:start_sample + chunk_samples]
        chunks.append(chunk)

    return chunks

# Sicherstellen, dass der Ausgabeordner existiert
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Neueste Audiodatei abrufen
latest_audio_file = get_latest_audio_file(input_folder)

if not latest_audio_file:
    print("Fehler: Keine Audiodateien im Ordner gefunden.")
else:
    try:
        print(f"Verarbeite Datei: {latest_audio_file}")
        audio_path = os.path.join(input_folder, latest_audio_file)

        # Teile die Audiodatei in 30 Sekunden Chunks
        audio_chunks = split_audio_into_chunks(audio_path, chunk_duration=30)

        # Word-Dokument erstellen
        doc = Document()
        doc.add_heading(f'Transkript der Datei: {latest_audio_file}', 0)

        # Für jedes Chunk transkribieren und in das Dokument einfügen
        for idx, chunk in enumerate(audio_chunks):
            print(f"Verarbeite Chunk {idx + 1} von {len(audio_chunks)}...")
            text, language = transcribe_audio_chunk(chunk)
            doc.add_paragraph(f"Erkannte Sprache im Chunk {idx + 1}: {language.upper()}\n")
            
            # Jeden Satz in eine eigene Zeile schreiben
            sentences = text.split(". ")
            for sentence in sentences:
                if sentence.strip():
                    doc.add_paragraph(sentence.strip())

        # Speichern der Datei
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(latest_audio_file)[0]}.docx")
        doc.save(output_file_path)

        print(f"Transkription erfolgreich gespeichert unter: {output_file_path}")

    except Exception as e:
        print(f"Ein Fehler ist aufgetreten: {e}")
