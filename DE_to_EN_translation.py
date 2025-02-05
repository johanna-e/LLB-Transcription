# Deutsch zu Englisch Übersetzung

import os
import whisper
from docx import Document
import torch

# Prüfen, ob eine GPU verfügbar ist
device = "cuda" if torch.cuda.is_available() else "cpu"
if device == "cuda":
    print(f"GPU wird verwendet: {torch.cuda.get_device_name(0)}")
else:
    print("Keine GPU verfügbar. Whisper läuft auf der CPU.")

# Funktion, um alle Audiodateien im Ordner zu finden
def get_audio_files(folder_path):
    audio_files = [f for f in os.listdir(folder_path) if f.endswith(('.mp3', '.wav', '.m4a'))]
    return audio_files

# Funktion, um eine Datei zu übersetzen
def translate_file(file_path, model):
    try:
        print(f"Übersetze Datei: {file_path}")
        result = model.transcribe(file_path, task="translate", language="de")  # Übersetzung ins Englische
        return result["text"]
    except Exception as e:
        print(f"Fehler bei der Übersetzung der Datei {file_path}: {e}")
        return None

# Funktion, um alle Audiodateien im Ordner zu übersetzen
def process_files(choice, input_folder, output_folder, model):
    if choice == 1:
        audio_files = get_audio_files(input_folder)
    elif choice == 2:
        audio_files = [get_latest_audio_file(input_folder)]
    elif choice == 3:
        audio_file_name = input("Bitte den Dateinamen (mit korrekter Endung, bspw.: .m4a) eingeben: ")
        audio_files = [audio_file_name]
    else:
        print("Ungültige Wahl!")
        return

    if not audio_files:
        print("Keine Audiodateien im Ordner gefunden.")
        return

    for audio_file in audio_files:
        file_path = os.path.join(input_folder, audio_file)

        if not os.path.exists(file_path):
            print(f"Die Datei {audio_file} wurde nicht gefunden.")
            continue

        print(f"Verarbeite Datei: {audio_file}")
        translated_text = translate_file(file_path, model)

        if translated_text:
            # Erstelle ein neues Word-Dokument
            output_file = os.path.join(output_folder, f"{os.path.splitext(audio_file)[0]}_translated.docx")
            doc = Document()
            doc.add_heading(f'Translation of: {audio_file}', 0)
            doc.add_paragraph(translated_text)
            doc.save(output_file)
            print(f"Übersetzung gespeichert unter: {output_file}")

# Funktion, um die neueste Audiodatei zu finden
def get_latest_audio_file(folder_path):
    audio_files = [f for f in os.listdir(folder_path) if f.endswith(('.mp3', '.wav', '.m4a'))]
    return max(audio_files, key=lambda f: os.path.getmtime(os.path.join(folder_path, f))) if audio_files else None

# Hauptteil des Skripts
def main():
    input_folder = "Dream recordings"
    output_folder = "Dream transcripts"
    
    # Sicherstellen, dass der Ausgabeordner existiert
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Lade das Whisper-Modell
    model = whisper.load_model("large")

    # Benutzerwahl
    print("Möchtest du:")
    print("1 - Alle Dateien im Ordner übersetzen")
    print("2 - Nur die neueste Datei übersetzen")
    print("3 - Eine bestimmte Datei übersetzen")
    choice = int(input("Bitte gib 1, 2 oder 3 ein: "))

    # Dateien verarbeiten
    process_files(choice, input_folder, output_folder, model)

if __name__ == "__main__":
    main()
