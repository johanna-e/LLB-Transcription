# Transcribe multiple English audio files

import os
import whisper
import torch
import subprocess
from docx import Document

# Function to convert WMA to MP3, since Whisper only supports MP3, WAV, and M4A
def convert_wma_to_mp3(input_path, output_path):
    subprocess.run(['ffmpeg', '-i', input_path, output_path])

# Function to find all audio files in the input-folder
def get_audio_files(folder_path):
    # Check input-folder for MP3, WAV, M4A, and WMA files
    audio_files = [f for f in os.listdir(folder_path) if f.endswith(('.mp3', '.wav', '.m4a', '.WMA'))]
    return audio_files

# Check if a GPU is available
device = "cuda" if torch.cuda.is_available() else "cpu"
if device == "cuda":
    print(f"Use GPU: {torch.cuda.get_device_name(0)}")
else:
    print("No GPU found. Whisper runs via CPU.")

# Folderpaths
input_folder = "Traumberichte_Audio"
output_folder = "Traumberichte_Transkript"

# Attepmt to get all audio files in the input-folder
try:
    audio_files = get_audio_files(input_folder)
    
    if audio_files:
        print(f"Found audio files: {audio_files}")
        
        # Load whisper model on the correct device
        model = whisper.load_model("turbo", device=device)
        
        # Process each audio file in the input-folder
        for audio_file in audio_files:
            print(f"Transcribe file: {audio_file}")
            
            # If input-file is a WMA-file, convert it to MP3 first
            if audio_file.endswith('.wma'):
                mp3_file = os.path.splitext(audio_file)[0] + '.mp3'
                input_path = os.path.join(input_folder, audio_file)
                output_path = os.path.join(input_folder, mp3_file)
                convert_wma_to_mp3(input_path, output_path)
                audio_file = mp3_file  # Now use the MP3-file
            
            # Transcribe the file
            audio_path = os.path.join(input_folder, audio_file)
            result = model.transcribe(audio_path, language="en")  # This is where the language is set to English
            
            # Create a new Word document
            doc = Document()
            doc.add_heading(f'Transcript of: {audio_file}', 0)
            
            # Add the transcribed text sentence by sentence
            sentences = result["text"].split(". ")  # Splitting the text by sentences
            for sentence in sentences:
                if sentence.strip():  # Only add non-empty sentences
                    doc.add_paragraph(sentence.strip())  # Adds the sentence as a new paragraph
            
            # Check if the output-folder exists, if not create it
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            
            # Save transcription in the output-folder
            output_file_path = os.path.join(output_folder, f"{os.path.splitext(audio_file)[0]}.docx")
            doc.save(output_file_path)
            
            print(f"Transcription successfully saved under: {output_file_path}")
    
    else:
        print("Error: No audio files were found in the input-folder.")
        
except Exception as e:
    print(f"An error occurred: {e}")
