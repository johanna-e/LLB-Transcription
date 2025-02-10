# Skript um gezielt spezifische englische Audiodateien zu transkribieren 


import os
import whisper
from docx import Document

# Ordnerpfade
input_folder = "Dream recordings"
output_folder = "Dream transcripts"

# Benutzerdefinierte Eingabe des Dateinamens
file_name = input("Bitte den Namen der Audiodatei (inklusive Dateiendung, z.B. 'audio.mp3') eingeben: ")

# Überprüfen, ob die Datei existiert
audio_path = os.path.join(input_folder, file_name)

if not os.path.isfile(audio_path):
    print(f"Fehler: Die Datei '{file_name}' wurde im Ordner '{input_folder}' nicht gefunden.")
else:
    try:
        print(f"Audiodatei gefunden: {file_name}")
        
        # Lade das Whisper-Modell
        model = whisper.load_model("turbo")
        
        # Transkribiere die Datei
        result = model.transcribe(audio_path, language="en")  # Sprache optional angeben für bessere Genauigkeit

        # Erstelle ein neues Word-Dokument
        doc = Document()
        doc.add_heading(f'Transkript der Datei: {file_name}', 0)

        # Füge jeden Satz in eine neue Zeile ein
        for segment in result["segments"]:
            doc.add_paragraph(segment["text"])
        
        # Sicherstellen, dass der Ausgabeordner existiert
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        
        # Speichere die Transkription im Ordner "Dream transcripts"
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}.docx")
        doc.save(output_file_path)
        
        print(f"Transkription erfolgreich gespeichert unter: {output_file_path}")
    
    except Exception as e:
        print(f"Ein Fehler ist aufgetreten: {e}")
