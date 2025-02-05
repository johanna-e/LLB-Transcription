# Dream Transcript for English Audio Files

import os
import whisper
from docx import Document

# Funktion, um die neueste Audiodatei im Ordner zu finden
def get_latest_audio_file(folder_path):
    # Durchsuche den Ordner nach mp3-, wav- und m4a-Dateien
    audio_files = [f for f in os.listdir(folder_path) if f.endswith(('.mp3', '.wav', '.m4a'))]
    
    # Wenn keine Audiodateien im Ordner sind, gib None zurück
    if not audio_files:
        return None
    
    # Finde die Datei mit dem neuesten Änderungsdatum
    latest_file = max(audio_files, key=lambda f: os.path.getmtime(os.path.join(folder_path, f)))
    return latest_file

# Ordnerpfade
input_folder = "Dream recordings"
output_folder = "Dream transcripts"

# Versuche, die neueste Audiodatei aus dem Ordner zu holen
try:
    latest_audio_file = get_latest_audio_file(input_folder)
    
    if latest_audio_file:
        print(f"Neueste Audiodatei gefunden: {latest_audio_file}")
        
        # Lade das Whisper-Modell
        model = whisper.load_model("turbo")
        
        # Transkribiere die Datei
        audio_path = os.path.join(input_folder, latest_audio_file)
        result = model.transcribe(audio_path, language="en") 

        # Erstelle ein neues Word-Dokument
        doc = Document()
        doc.add_heading(f'Transkript der Datei: {latest_audio_file}', 0)

        # Füge jeden Satz in eine neue Zeile ein
        for segment in result["segments"]:
            doc.add_paragraph(segment["text"])
        
        # Sicherstellen, dass der Ausgabeordner existiert
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        
        # Speichere die Transkription im Ordner "Dream transcripts"
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(latest_audio_file)[0]}.docx")
        doc.save(output_file_path)
        
        print(f"Transkription erfolgreich gespeichert unter: {output_file_path}")
    
    else:
        print("Fehler: Keine Audiodateien im Ordner 'Dream recordings' gefunden.")
        
except Exception as e:
    print(f"Ein Fehler ist aufgetreten: {e}")
